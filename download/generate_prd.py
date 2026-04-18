#!/usr/bin/env python3
"""
Generate PRD KNBMP PGA-72 CMS Document
Professional Product Requirements Document in DOCX format
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

# ── Color Constants ──────────────────────────────────────────
BURGUNDY = RGBColor(0x5E, 0x21, 0x29)
GOLD = RGBColor(0xC5, 0xA0, 0x59)
PARCHMENT = RGBColor(0xF5, 0xF1, 0xE8)
DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D)
MEDIUM_TEXT = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GOLD = RGBColor(0xF5, 0xEE, 0xD5)
TABLE_HEADER_BG = RGBColor(0x5E, 0x21, 0x29)
TABLE_ALT_BG = RGBColor(0xFD, 0xF8, 0xF0)

# Domain colors
DOMAIN_COLORS = {
    "D1": "C4952A", "D2": "1565C0", "D3": "6A1B9A", "D4": "008F3D",
    "D5": "E65100", "D6": "AD1457", "D7": "00838F", "D8": "4527A0", "D9": "BF360C"
}

# ── Helper Functions ─────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{val.get("sz", 4)}" '
                f'w:space="0" w:color="{val.get("color", "5E2129")}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_paragraph_spacing(paragraph, before=0, after=0, line=312):
    """Set paragraph spacing with 1.3x line spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line / 20)  # 312 twips ≈ 15.6pt ≈ 1.3x at 12pt

def add_formatted_run(paragraph, text, bold=False, italic=False, size=11, color=DARK_TEXT, font_name="Calibri"):
    """Add a formatted text run to paragraph."""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    return run

def set_cell_text(cell, text, bold=False, size=10, color=DARK_TEXT, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name="Calibri"):
    """Set text in a table cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    set_paragraph_spacing(p, before=2, after=2, line=276)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font_name
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return run

def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a professional styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Set table margins
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    margins = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="40" w:type="dxa"/>'
        f'<w:left w:w="80" w:type="dxa"/>'
        f'<w:bottom w:w="40" w:type="dxa"/>'
        f'<w:right w:w="80" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(margins)
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "5E2129")
        set_cell_text(cell, header, bold=True, size=10, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, "FDF8F0")
            text = str(cell_text)
            is_first = c_idx == 0
            set_cell_text(cell, text, bold=is_first, size=9, color=DARK_TEXT)
    
    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)
    
    return table

def add_bullet_point(doc, text, level=0, bold_prefix=""):
    """Add a bullet point with optional bold prefix."""
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    set_paragraph_spacing(p, before=2, after=2, line=276)
    if bold_prefix:
        add_formatted_run(p, bold_prefix, bold=True, size=10)
        add_formatted_run(p, text, size=10)
    else:
        add_formatted_run(p, text, size=10)
    return p

def add_numbered_item(doc, number, text, bold_prefix=""):
    """Add a numbered item."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    set_paragraph_spacing(p, before=2, after=2, line=276)
    add_formatted_run(p, f"{number}. ", bold=True, size=10, color=BURGUNDY)
    if bold_prefix:
        add_formatted_run(p, bold_prefix, bold=True, size=10)
        add_formatted_run(p, text, size=10)
    else:
        add_formatted_run(p, text, size=10)
    return p

def add_section_heading(doc, text, level=1):
    """Add a section heading with proper formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BURGUNDY if level <= 2 else DARK_TEXT
        run.font.name = "DM Serif Display" if level == 1 else "Calibri"
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    set_paragraph_spacing(h, before=18 if level == 1 else 12, after=8, line=312)
    return h

def add_body_paragraph(doc, text, bold=False, italic=False):
    """Add a body paragraph with standard formatting."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=3, after=6, line=312)
    add_formatted_run(p, text, bold=bold, italic=italic, size=11)
    return p

def add_separator(doc):
    """Add a gold decorative line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=6, line=312)
    run = p.add_run("━" * 60)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)
    return p

def add_info_box(doc, title, content, bg_color="F5F1E8"):
    """Add a highlighted information box using a table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    
    # Title
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=4, after=2, line=276)
    add_formatted_run(p, title, bold=True, size=10, color=BURGUNDY)
    
    # Content
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=2, after=4, line=276)
    add_formatted_run(p2, content, size=9, color=DARK_TEXT)
    
    # Border
    set_cell_borders(cell, 
        bottom={"sz": 8, "color": "C5A059"},
        left={"sz": 12, "color": "5E2129"})
    
    doc.add_paragraph()  # spacing
    return table


# ── Main Document Creation ───────────────────────────────────

def create_prd():
    doc = Document()
    
    # ── Page Setup (A4) ──
    for section in doc.sections:
        section.page_width = Twips(11906)
        section.page_height = Twips(16838)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # ── Define Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT
    pf = style.paragraph_format
    pf.line_spacing = Pt(14.3)  # 1.3x at 11pt
    
    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.color.rgb = BURGUNDY
    h1_style.font.name = 'DM Serif Display'
    h1_style.font.size = Pt(18)
    
    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.color.rgb = BURGUNDY
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    
    # Heading 3 style
    h3_style = doc.styles['Heading 3']
    h3_style.font.color.rgb = DARK_TEXT
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(12)

    # ═══════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════
    
    # Top spacer
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
    
    # Gold decorative line
    add_separator(doc)
    
    # Organization name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=6, line=312)
    add_formatted_run(p, "KOPERASI NASIONAL MULTIPEDO (KNMP)", bold=True, size=12, color=GOLD, font_name="Calibri")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6, line=312)
    add_formatted_run(p, "Jaringan Ekosistem — Paguyuban, Pemuda, Perempuan (JE-P3)", bold=False, size=10, color=MEDIUM_TEXT)
    
    # Spacer
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
    
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=0, line=312)
    add_formatted_run(p, "PRD", bold=True, size=42, color=BURGUNDY, font_name="DM Serif Display")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line=312)
    add_formatted_run(p, "━" * 30, size=14, color=GOLD)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4, line=360)
    add_formatted_run(p, "KNBMP PGA-72", bold=True, size=28, color=BURGUNDY, font_name="DM Serif Display")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2, line=360)
    add_formatted_run(p, "Content Management System", bold=True, size=22, color=GOLD, font_name="DM Serif Display")
    
    # Spacer
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=6, line=312)
    add_formatted_run(p, "Product Requirements Document", italic=True, size=14, color=MEDIUM_TEXT)
    add_formatted_run(p, " — ", size=14, color=GOLD)
    add_formatted_run(p, "Dari Hulu ke Hilir", italic=True, size=14, color=MEDIUM_TEXT)
    
    # Version & Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=2, line=312)
    add_formatted_run(p, f"v1.0 — Dokumen Hidup", bold=True, size=12, color=BURGUNDY)
    
    today = datetime.date.today()
    date_str = today.strftime("%d %B %Y")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=6, line=312)
    add_formatted_run(p, date_str, size=11, color=MEDIUM_TEXT)
    
    # Spacer
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
    
    add_separator(doc)
    
    # Confidential
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=2, line=312)
    add_formatted_run(p, "RAHASIA — CONFIDENTIAL", bold=True, size=10, color=BURGUNDY)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=2, line=276)
    add_formatted_run(p, "Dokumen ini bersifat rahasia dan hanya boleh digunakan oleh pihak yang berwenang.", italic=True, size=8, color=MEDIUM_TEXT)
    
    # Page break after cover
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # DAFTAR ISI (TABLE OF CONTENTS)
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "DAFTAR ISI", level=1)
    add_separator(doc)
    doc.add_paragraph()
    
    toc_items = [
        ("BAB 1", "Pendahuluan", ""),
        ("   1.1", "Latar Belakang", ""),
        ("   1.2", "Tujuan Dokumen", ""),
        ("   1.3", "Ruang Lingkup", ""),
        ("   1.4", "Definisi & Istilah", ""),
        ("   1.5", "Pemangku Kepentingan", ""),
        ("BAB 2", "Visi Produk", ""),
        ("   2.1", "Visi CMS", ""),
        ("   2.2", "Target Pengguna", ""),
        ("   2.3", "User Persona", ""),
        ("   2.4", "Alur Kerja Konten", ""),
        ("BAB 3", "Arsitektur Sistem", ""),
        ("   3.1", "Arsitektur High-Level", ""),
        ("   3.2", "Tech Stack", ""),
        ("   3.3", "Database Schema", ""),
        ("   3.4", "API Endpoints", ""),
        ("   3.5", "Authentication", ""),
        ("BAB 4", "Fitur CMS (Admin Panel)", ""),
        ("   4.1", "Dashboard", ""),
        ("   4.2", "Manajemen Halaman", ""),
        ("   4.3", "Manajemen PGA-72", ""),
        ("   4.4", "Editor Konten", ""),
        ("   4.5", "Media Manager", ""),
        ("   4.6", "Theme & Design", ""),
        ("   4.7", "Pengaturan Situs", ""),
        ("   4.8", "Audit Log", ""),
        ("BAB 5", "Struktur Konten", ""),
        ("   5.1", "Peta Halaman", ""),
        ("   5.2", "9 Domain Detail", ""),
        ("   5.3", "Struktur Data Per PGA", ""),
        ("BAB 6", "Desain Antarmuka", ""),
        ("   6.1", "Design System", ""),
        ("   6.2", "Typography", ""),
        ("   6.3", "Batik Kawung Pattern", ""),
        ("   6.4", "Responsive Design", ""),
        ("   6.5", "Flipbook Interactions", ""),
        ("BAB 7", "Implementasi", ""),
        ("   7.1", "Roadmap 3 Fase", ""),
        ("   7.2", "Deployment", ""),
        ("   7.3", "Backup & Recovery", ""),
        ("   7.4", "Monitoring", ""),
        ("BAB 8", "Keamanan", ""),
        ("   8.1", "Autentikasi", ""),
        ("   8.2", "Otorisasi", ""),
        ("   8.3", "Validasi Input", ""),
        ("   8.4", "Rate Limiting", ""),
        ("   8.5", "Audit Trail", ""),
        ("BAB 9", "Risiko & Mitigasi", ""),
        ("BAB 10", "Lampiran", ""),
        ("   10.1", "API Reference", ""),
        ("   10.2", "Database ERD", ""),
        ("   10.3", "Skema Prisma", ""),
    ]
    
    for num, title, _ in toc_items:
        p = doc.add_paragraph()
        is_chapter = num.strip().startswith("BAB")
        p.paragraph_format.left_indent = Inches(0) if is_chapter else Inches(0.5)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=2, leader=1)  # Right-aligned with dot leader
        set_paragraph_spacing(p, before=1 if not is_chapter else 4, after=1, line=276)
        add_formatted_run(p, f"{num}  ", bold=is_chapter, size=11 if is_chapter else 10, 
                         color=BURGUNDY if is_chapter else DARK_TEXT)
        add_formatted_run(p, title, bold=is_chapter, size=11 if is_chapter else 10,
                         color=BURGUNDY if is_chapter else DARK_TEXT)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=6, line=276)
    add_formatted_run(p, "Catatan: Daftar isi ini akan diperbarui secara otomatis saat dokumen dibuka di Microsoft Word.", italic=True, size=9, color=MEDIUM_TEXT)
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 1: PENDAHULUAN
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 1: PENDAHULUAN", level=1)
    add_separator(doc)
    
    # 1.1 Latar Belakang
    add_section_heading(doc, "1.1 Latar Belakang", level=2)
    
    add_body_paragraph(doc, 
        "Koperasi Nasional Multipedo (KNMP) adalah gerakan koperasi yang bertujuan mewujudkan kedaulatan ekonomi "
        "untuk 270 juta rakyat Indonesia yang tersebar di 83.763 desa di seluruh Nusantara. Sebagai bentuk "
        "komitmen terhadap transparansi dan profesionalisme, KNMP telah mengembangkan aplikasi flipbook mewah "
        "berbasis Next.js 16 yang menampilkan dokumen komprehensif 93 halaman berisi 72 Pilar Fondasional (PGA).")
    
    add_body_paragraph(doc, 
        "Namun, dalam proses pengembangan sistem ini, terjadi kerugian finansial yang sangat signifikan sebesar "
        "USD 1.000.000 (satu juta dolar Amerika Serikat). Kerugian ini disebabkan oleh berbagai faktor termasuk "
        "ketergantungan pada pihak ketiga yang tidak dapat dipercaya, kurangnya dokumentasi teknis yang memadai, "
        "serta ketiadaan sistem manajemen konten yang memungkinkan pengelolaan mandiri oleh tim internal KNMP.")
    
    add_body_paragraph(doc,
        "Saat ini, pemangku kepentingan KNMP tidak memiliki kemampuan teknis (coding) untuk melakukan perubahan "
        "konten pada flipbook. Setiap perubahan, bahkan yang terkecil sekalipun, memerlukan keterlibatan developer "
        "eksternal yang memakan waktu dan biaya. Kondisi ini tidak berkelanjutan dan bertentangan dengan prinsip "
        "kedaulatan yang menjadi fondasi gerakan koperasi ini.")
    
    add_info_box(doc, "MENGAPA CMS INI PENTING?",
        "Dengan CMS yang tepat, tim KNMP dapat mengelola konten flipbook secara mandiri tanpa perlu menulis kode "
        "satu baris pun. Perubahan konten di backend akan otomatis ter-update di frontend, menghemat waktu dan biaya "
        "secara signifikan. Ini adalah langkah krusial menuju kemandirian digital.")
    
    add_body_paragraph(doc,
        "Dokumen PRD ini disusun sebagai panduan lengkap untuk membangun Content Management System (CMS) yang "
        "berfungsi sebagai \"WordPress-nya KNMP\" — sebuah platform yang memberdayakan pengguna non-teknis untuk "
        "mengelola seluruh konten flipbook PGA-72 dari ujung ke ujung, dari hulu (data input) hingga hilir "
        "(tampilan di flipbook).")
    
    # 1.2 Tujuan Dokumen
    add_section_heading(doc, "1.2 Tujuan Dokumen", level=2)
    
    add_body_paragraph(doc, "Dokumen PRD ini memiliki tujuan-tujuan berikut:")
    
    objectives = [
        ("Memberikan panduan lengkap ", "untuk membangun CMS KNBMP PGA-72 dari nol hingga siap digunakan."),
        ("Mendefinisikan arsitektur sistem ", "yang menghubungkan Frontend (Flipbook), API Layer, Database, dan Admin Panel secara terintegrasi."),
        ("Menetapkan standar keamanan ", "yang memadai untuk melindungi data sensitif organisasi dan 270 juta anggota potensial."),
        ("Membuat roadmap implementasi ", "yang realistis dalam 3 fase selama 6 minggu."),
        ("Memastikan skalabilitas ", "sistem untuk mengakomodasi pertumbuhan konten dan pengguna di masa depan."),
        ("Menjadi dokumen referensi tunggal (single source of truth) ", "bagi seluruh tim pengembang dan pemangku kepentingan."),
    ]
    
    for i, (bold_part, normal_part) in enumerate(objectives, 1):
        add_numbered_item(doc, i, normal_part, bold_part)
    
    doc.add_paragraph()
    
    # 1.3 Ruang Lingkup
    add_section_heading(doc, "1.3 Ruang Lingkup", level=2)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Dalam Ruang Lingkup (In-Scope):", bold=True, size=11, color=BURGUNDY)
    
    in_scope = [
        "Admin Panel (Dashboard) untuk manajemen konten flipbook",
        "CRUD (Create, Read, Update, Delete) untuk seluruh 93+ halaman flipbook",
        "Editor konten khusus untuk 72 Pilar Fondasional (PGA)",
        "Sistem autentikasi dan otorisasi berbasis peran (RBAC)",
        "Media Manager untuk upload dan pengelolaan gambar/dokumen",
        "API RESTful untuk komunikasi antara Admin Panel dan Flipbook",
        "Database terpusat untuk menyimpan seluruh konten",
        "Audit log untuk melacak setiap perubahan konten",
        "Tema dan desain yang konsisten dengan identitas visual KNMP",
        "Auto-update frontend ketika konten berubah di backend",
    ]
    for item in in_scope:
        add_bullet_point(doc, item)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Di Luar Ruang Lingkup (Out-of-Scope):", bold=True, size=11, color=BURGUNDY)
    
    out_scope = [
        "Aplikasi mobile native (iOS/Android)",
        "Sistem pembayaran atau e-commerce",
        "Integrasi dengan media sosial secara real-time",
        "Multi-bahasa (hanya Bahasa Indonesia)",
        "Hosting dan infrastruktur cloud (akan ditangani terpisah)",
        "Migrasi data dari sistem lama (jika ada)",
    ]
    for item in out_scope:
        add_bullet_point(doc, item)
    
    # 1.4 Definisi & Istilah
    add_section_heading(doc, "1.4 Definisi & Istilah", level=2)
    
    glossary_data = [
        ["PGA", "Pilar Fondasional", "72 pilar utama yang menjadi fondasi kerangka kerja KNBMP"],
        ["KPA", "Kerangka Pemikiran Awal", "Dokumen awal yang menjadi dasar penyusunan PGA"],
        ["KND", "Kerangka Nasional Digital", "Kerangka digital nasional Indonesia yang menjadi referensi"],
        ["LKD", "Laporan Kerja Digital", "Dokumen laporan pelaksanaan kerja secara digital"],
        ["KNMP", "Koperasi Nasional Multipedo", "Organisasi induk yang menaungi gerakan koperasi"],
        ["KNBMP", "Koperasi Nasional Bersama Multipedo Persada", "Entitas operasional utama"],
        ["JE-P3", "Jaringan Ekosistem Paguyuban, Pemuda, Perempuan", "Network ekosistem KNMP"],
        ["CMS", "Content Management System", "Sistem manajemen konten"],
        ["Flipbook", "Aplikasi pembaca buku digital", "Aplikasi Next.js 16 untuk menampilkan dokumen PGA-72"],
        ["RBAC", "Role-Based Access Control", "Sistem kontrol akses berdasarkan peran pengguna"],
        ["API", "Application Programming Interface", "Antarmuka pemrograman aplikasi"],
        ["CRUD", "Create, Read, Update, Delete", "Operasi dasar pengelolaan data"],
        ["MDX", "Markdown + JSX", "Format konten yang mendukung Markdown dan komponen React"],
        ["Domain", "Area fungsional", "9 area fungsional yang mengelompokkan 72 PGA"],
        ["Prisma", "ORM untuk Node.js", "Object-Relational Mapping untuk database"],
        ["NextAuth.js", "Library autentikasi Next.js", "Library untuk implementasi autentikasi"],
    ]
    
    create_styled_table(doc, ["Istilah", "Nama Lengkap", "Deskripsi"], glossary_data, [1.0, 1.8, 3.5])
    
    doc.add_paragraph()
    
    # 1.5 Pemangku Kepentingan
    add_section_heading(doc, "1.5 Pemangku Kepentingan", level=2)
    
    stakeholder_data = [
        ["Ketua Umum KNMP", "Pemimpin Tertinggi", "Pemberi persetujuan akhir, sponsor utama"],
        ["Sekretaris Jenderal", "Pengawas Operasional", "Pengawasan implementasi dan pelaporan"],
        ["Tim Konten", "Editor", "Pengguna utama CMS, pengelola konten harian"],
        ["Tim IT Internal", "Administrator", "Pengelola teknis sistem, maintenance"],
        ["Developer Eksternal", "Implementor", "Pengembang yang membangun CMS"],
        ["Pengawas Keuangan", "Reviewer", "Pengawasan anggaran pengembangan"],
        ["Anggota Koperasi", "End User", "Pembaca flipbook, penerima manfaat"],
    ]
    
    create_styled_table(doc, ["Pemangku Kepentingan", "Peran", "Deskripsi"], stakeholder_data, [1.5, 1.5, 3.3])
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 2: VISI PRODUK
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 2: VISI PRODUK", level=1)
    add_separator(doc)
    
    # 2.1 Visi CMS
    add_section_heading(doc, "2.1 Visi CMS", level=2)
    
    add_body_paragraph(doc,
        "Visi utama dari sistem CMS ini adalah menjadi \"WordPress-nya KNMP\" — sebuah platform manajemen konten "
        "yang memberdayakan setiap anggota tim KNMP untuk mengelola, memperbarui, dan mempublikasikan konten "
        "flipbook PGA-72 tanpa memerlukan keahlian teknis apapun.")
    
    add_info_box(doc, "VISI",
        "\"Menciptakan sistem manajemen konten yang memberdayakan 83.763 desa dan 270 juta rakyat Indonesia "
        "untuk mengelola informasi strategis koperasi secara mandiri, transparan, dan efisien — dari desa "
        "hingga pusat, dari hulu hingga hilir.\"")
    
    add_body_paragraph(doc, "CMS ini dirancang dengan prinsip-prinsip berikut:")
    
    principles = [
        ("Kemudahan Penggunaan (Usability): ", "Antarmuka yang intuitif, bahkan untuk pengguna yang tidak memiliki latar belakang teknis."),
        ("Kemandirian (Autonomy): ", "Tim internal KNMP dapat sepenuhnya mengelola konten tanpa ketergantungan pada developer."),
        ("Keamanan (Security): ", "Lapisan keamanan berlapis untuk melindungi data sensitif organisasi."),
        ("Skalabilitas (Scalability): ", "Sistem dapat berkembang seiring pertumbuhan konten dan pengguna."),
        ("Transparansi (Transparency): ", "Setiap perubahan tercatat dalam audit log untuk akuntabilitas."),
    ]
    for bold_part, normal_part in principles:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 2.2 Target Pengguna
    add_section_heading(doc, "2.2 Target Pengguna", level=2)
    
    add_body_paragraph(doc, "CMS ini dirancang untuk tiga kategori utama pengguna:")
    
    user_types = [
        ("Super Admin", "Akses penuh ke seluruh fitur CMS, manajemen pengguna, pengaturan sistem, dan audit log. Hanya 1-2 orang yang memegang peran ini."),
        ("Editor", "Akses untuk membuat, mengedit, dan mengelola konten halaman dan PGA. Dapat mempublikasikan kontan setelah review. Tim konten utama."),
        ("Reviewer", "Akses untuk meninjau dan menyetujui konten sebelum dipublikasikan. Tidak dapat mengedit langsung, hanya memberikan persetujuan."),
    ]
    
    for bold_part, normal_part in user_types:
        add_bullet_point(doc, normal_part, bold_prefix=f"{bold_part}: ")
    
    # 2.3 User Persona
    add_section_heading(doc, "2.3 User Persona", level=2)
    
    # Persona 1
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=312)
    add_formatted_run(p, "Persona 1: Bu Siti Rahmawati — Sekretaris Konten", bold=True, size=12, color=BURGUNDY)
    
    persona1_data = [
        ["Usia", "45 tahun"],
        ["Pendidikan", "S1 Manajemen"],
        ["Pengalaman Teknis", "Dasar Microsoft Office, tidak bisa coding"],
        ["Peran di CMS", "Editor"],
        ["Tugas Utama", "Mengelola konten 72 PGA, memperbarui informasi halaman"],
        ["Pain Point", "Setiap perubahan harus menunggu developer, memakan waktu berminggu-minggu"],
        ["Expectation", "Dapat mengubah konten sendiri dan melihat hasilnya langsung di flipbook"],
    ]
    create_styled_table(doc, ["Atribut", "Detail"], persona1_data, [1.5, 4.8])
    doc.add_paragraph()
    
    # Persona 2
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=312)
    add_formatted_run(p, "Persona 2: Pak Ahmad Fauzi — Ketua Tim Pengawas", bold=True, size=12, color=BURGUNDY)
    
    persona2_data = [
        ["Usia", "52 tahun"],
        ["Pendidikan", "S2 Hukum"],
        ["Pengalaman Teknis", "Pengguna email dan WhatsApp, tidak familiar dengan sistem web"],
        ["Peran di CMS", "Reviewer"],
        ["Tugas Utama", "Meninjau dan menyetujui perubahan konten sebelum publikasi"],
        ["Pain Point", "Tidak bisa memantau perubahan konten secara real-time"],
        ["Expectation", "Dashboard sederhana yang menampilkan perubahan pending dan riwayat revisi"],
    ]
    create_styled_table(doc, ["Atribut", "Detail"], persona2_data, [1.5, 4.8])
    doc.add_paragraph()
    
    # Persona 3
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=312)
    add_formatted_run(p, "Persona 3: Mas Dimas Pratama — Admin IT", bold=True, size=12, color=BURGUNDY)
    
    persona3_data = [
        ["Usia", "28 tahun"],
        ["Pendidikan", "S1 Teknik Informatika"],
        ["Pengalaman Teknis", "Familiar dengan Linux, Docker, dan dasar-dasar web development"],
        ["Peran di CMS", "Super Admin"],
        ["Tugas Utama", "Mengelola pengguna, pengaturan sistem, backup data, dan monitoring"],
        ["Pain Point", "Perlu memahami arsitektur sistem dengan cepat untuk melakukan maintenance"],
        ["Expectation", "Dokumentasi teknis lengkap dan dashboard monitoring yang informatif"],
    ]
    create_styled_table(doc, ["Atribut", "Detail"], persona3_data, [1.5, 4.8])
    
    # 2.4 Alur Kerja Konten
    add_section_heading(doc, "2.4 Alur Kerja Konten", level=2)
    
    add_body_paragraph(doc,
        "Alur kerja konten CMS KNBMP mengikuti pola Draft → Review → Publish untuk memastikan kualitas "
        "dan akurasi setiap informasi yang dipublikasikan melalui flipbook.")
    
    workflow_steps = [
        ["1. Draft", "Editor membuat atau mengedit konten melalui Editor Konten CMS. Konten disimpan sebagai draft dan belum terlihat di flipbook."],
        ["2. Preview", "Editor dapat melihat preview tampilan konten di flipbook sebelum mengirimkan untuk review."],
        ["3. Submit Review", "Editor mengirimkan draft ke Reviewer untuk mendapatkan persetujuan. Notifikasi dikirim ke Reviewer."],
        ["4. Review", "Reviewer memeriksa konten, memberikan catatan revisi jika diperlukan, atau menyetujui."],
        ["5. Revisi (jika perlu)", "Jika ada catatan revisi, Editor melakukan perubahan dan mengirimkan kembali untuk review."],
        ["6. Approve", "Reviewer menyetujui konten. Status berubah dari 'pending' menjadi 'approved'."],
        ["7. Publish", "Konten yang disetujui dipublikasikan dan otomatis muncul di flipbook frontend."],
        ["8. Audit", "Seluruh proses tercatat dalam audit log termasuk siapa, kapan, dan perubahan apa yang dilakukan."],
    ]
    
    create_styled_table(doc, ["Tahap", "Deskripsi"], workflow_steps, [1.3, 5.0])
    
    add_info_box(doc, "AUTO-UPDATE MEKANISME",
        "Setiap kali konten dipublish melalui CMS, sistem akan secara otomatis men-generate ulang data "
        "flipbook dan mengirimkan notifikasi ke frontend. Pengunjung flipbook akan melihat konten terbaru "
        "tanpa perlu me-refresh halaman (real-time update via WebSocket atau polling).")
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 3: ARSITEKTUR SISTEM
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 3: ARSITEKTUR SISTEM", level=1)
    add_separator(doc)
    
    # 3.1 Arsitektur High-Level
    add_section_heading(doc, "3.1 Arsitektur High-Level", level=2)
    
    add_body_paragraph(doc,
        "Arsitektur sistem CMS KNBMP PGA-72 mengikuti pola modern full-stack monorepo dengan "
        "pemisahan yang jelas antara frontend (Flipbook), API layer, database, dan admin panel. "
        "Semua komponen berada dalam satu proyek Next.js 16 yang terintegrasi.")
    
    add_body_paragraph(doc, "Komponen utama arsitektur:")
    
    arch_components = [
        ("Frontend Flipbook (Public): ", "Aplikasi Next.js 16 yang menampilkan konten flipbook kepada publik. Menggunakan Server Components untuk performa optimal dan client-side interactivity untuk animasi page flip."),
        ("Admin Panel (Protected): ", "Antarmuka manajemen konten yang dilindungi autentikasi. Dibangun menggunakan shadcn/ui untuk komponen UI yang konsisten dan dapat diakses."),
        ("API Routes (Next.js): ", "RESTful API yang menangani komunikasi antara frontend, admin panel, dan database. Menggunakan Next.js Route Handlers."),
        ("Database (SQLite + Prisma): ", "Database relasional ringan menggunakan SQLite dengan Prisma ORM untuk type-safe database access. Data disimpan dalam file database tunggal untuk kemudahan backup."),
        ("Auth Layer (NextAuth.js): ", "Sistem autentikasi dan otorisasi menggunakan NextAuth.js v5 dengan strategi credentials provider untuk login berbasis email dan password."),
    ]
    for bold_part, normal_part in arch_components:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    doc.add_paragraph()
    add_body_paragraph(doc, "Alur Data Sistem:")
    
    flow_items = [
        "Admin login → NextAuth.js verifikasi kredensial → Session token diterbitkan",
        "Admin edit konten via Admin Panel → Validasi Zod → Simpan ke SQLite via Prisma",
        "Data tersimpan → API generate static params → Flipbook di-revalidate (ISR)",
        "Pengunjung akses flipbook → Server Component fetch data dari database → Render halaman",
    ]
    for item in flow_items:
        add_bullet_point(doc, item)
    
    # 3.2 Tech Stack
    add_section_heading(doc, "3.2 Tech Stack", level=2)
    
    tech_stack = [
        ["Framework", "Next.js 16", "Full-stack React framework dengan App Router, Server Components, dan ISR"],
        ["Bahasa", "TypeScript 5.x", "Type-safe JavaScript untuk reliability kode"],
        ["UI Library", "shadcn/ui + Tailwind CSS 4", "Komponen UI yang customizable dan modern"],
        ["Styling", "Tailwind CSS 4", "Utility-first CSS framework"],
        ["Database", "SQLite 3", "Database relasional ringan, file-based"],
        ["ORM", "Prisma 6.x", "Type-safe ORM untuk database access"],
        ["Autentikasi", "NextAuth.js v5", "Authentication library untuk Next.js"],
        ["Validasi", "Zod", "Schema validation untuk input data"],
        ["Editor", "MDX / Rich Text", "Editor konten mendukung Markdown dan formatting kaya"],
        ["Font", "DM Serif Display, Inter, Libre Baskerville", "Typography heritage theme"],
        ["Deployment", "Vercel / Standalone", "Cloud deployment atau self-hosted"],
        ["Package Manager", "Bun", "Fast JavaScript package manager"],
    ]
    
    create_styled_table(doc, ["Kategori", "Teknologi", "Deskripsi"], tech_stack, [1.2, 2.0, 3.1])
    
    # 3.3 Database Schema
    add_section_heading(doc, "3.3 Database Schema", level=2)
    
    add_body_paragraph(doc,
        "Database menggunakan SQLite dengan Prisma ORM. Berikut adalah tabel-tabel utama dalam sistem:")
    
    add_section_heading(doc, "Tabel: users", level=3)
    users_schema = [
        ["id", "String (UUID)", "Primary key, auto-generated"],
        ["name", "String", "Nama lengkap pengguna"],
        ["email", "String (unique)", "Email untuk login"],
        ["password", "String (hashed)", "Password yang di-hash menggunakan bcrypt"],
        ["role", "Enum (SUPER_ADMIN, EDITOR, REVIEWER)", "Peran pengguna dalam sistem"],
        ["avatar", "String (optional)", "URL avatar pengguna"],
        ["isActive", "Boolean", "Status akun (active/suspended)"],
        ["lastLogin", "DateTime (optional)", "Waktu login terakhir"],
        ["createdAt", "DateTime", "Waktu pembuatan akun"],
        ["updatedAt", "DateTime", "Waktu update terakhir"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], users_schema, [1.2, 2.0, 3.1])
    
    doc.add_paragraph()
    add_section_heading(doc, "Tabel: pages", level=3)
    pages_schema = [
        ["id", "String (UUID)", "Primary key"],
        ["pageNumber", "Int (unique)", "Nomor halaman dalam flipbook (1-96)"],
        ["type", "Enum", "COVER, FOREWORD, MUKADIMAH, TOC, PGA, PHILOSOPHY, PACT, BACK_COVER"],
        ["title", "String", "Judul halaman"],
        ["slug", "String (unique)", "URL-friendly identifier"],
        ["content", "Text (MDX)", "Konten halaman dalam format MDX"],
        ["status", "Enum", "DRAFT, PENDING_REVIEW, APPROVED, PUBLISHED"],
        ["domainId", "String (FK)", "Referensi ke domain (untuk halaman PGA)"],
        ["authorId", "String (FK)", "Referensi ke user yang membuat"],
        ["reviewerId", "String (FK, optional)", "Referensi ke reviewer yang menyetujui"],
        ["publishedAt", "DateTime (optional)", "Waktu publikasi"],
        ["createdAt", "DateTime", "Waktu pembuatan"],
        ["updatedAt", "DateTime", "Waktu update terakhir"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], pages_schema, [1.2, 2.0, 3.1])
    
    doc.add_paragraph()
    add_section_heading(doc, "Tabel: pga_pillars", level=3)
    pga_schema = [
        ["id", "String (UUID)", "Primary key"],
        ["code", "String (unique)", "Kode PGA (PGA-01 s/d PGA-72)"],
        ["name", "String", "Nama PGA dalam Bahasa Indonesia"],
        ["nameEn", "String", "Nama PGA dalam Bahasa Inggris"],
        ["description", "Text", "Deskripsi lengkap PGA"],
        ["vision", "String", "Vision statement PGA"],
        ["domainId", "String (FK)", "Referensi ke domain"],
        ["orderInDomain", "Int", "Urutan dalam domain"],
        ["color", "String", "Warna representasi PGA (hex)"],
        ["dimensions", "JSON", "Array dimensi PGA"],
        ["principles", "JSON", "Array prinsip PGA"],
        ["crossReferences", "JSON", "Array referensi silang ke PGA lain"],
        ["metadata", "JSON", "Metadata tambahan (flexible)"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], pga_schema, [1.3, 1.8, 3.2])
    
    doc.add_paragraph()
    add_section_heading(doc, "Tabel: domains", level=3)
    domains_schema = [
        ["id", "String (UUID)", "Primary key"],
        ["code", "String (unique)", "Kode domain (D1 s/d D9)"],
        ["name", "String", "Nama domain dalam Bahasa Indonesia"],
        ["nameEn", "String", "Nama domain dalam Bahasa Inggris"],
        ["description", "Text", "Deskripsi domain"],
        ["color", "String", "Warna domain (hex)"],
        ["icon", "String (optional)", "Nama ikon/emoji representasi"],
        ["order", "Int", "Urutan tampilan"],
        ["pgaCount", "Int", "Jumlah PGA dalam domain (selalu 8)"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], domains_schema, [1.2, 1.8, 3.3])
    
    doc.add_paragraph()
    add_section_heading(doc, "Tabel: media", level=3)
    media_schema = [
        ["id", "String (UUID)", "Primary key"],
        ["filename", "String", "Nama file asli"],
        ["filepath", "String", "Path file di storage"],
        ["mimeType", "String", "Tipe MIME (image/jpeg, dll)"],
        ["size", "Int", "Ukuran file dalam bytes"],
        ["alt", "String (optional)", "Alt text untuk gambar"],
        ["uploadedById", "String (FK)", "User yang mengupload"],
        ["createdAt", "DateTime", "Waktu upload"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], media_schema, [1.3, 1.8, 3.2])
    
    doc.add_paragraph()
    add_section_heading(doc, "Tabel: settings", level=3)
    settings_schema = [
        ["id", "String (UUID)", "Primary key"],
        ["key", "String (unique)", "Kunci pengaturan"],
        ["value", "Text", "Nilai pengaturan"],
        ["type", "Enum", "STRING, NUMBER, BOOLEAN, JSON"],
        ["description", "String", "Deskripsi pengaturan"],
        ["updatedAt", "DateTime", "Waktu update terakhir"],
        ["updatedById", "String (FK)", "User yang mengubah"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], settings_schema, [1.3, 1.8, 3.2])
    
    doc.add_paragraph()
    add_section_heading(doc, "Tabel: audit_log", level=3)
    audit_schema = [
        ["id", "String (UUID)", "Primary key"],
        ["userId", "String (FK)", "User yang melakukan aksi"],
        ["action", "Enum", "CREATE, UPDATE, DELETE, PUBLISH, LOGIN, LOGOUT"],
        ["entityType", "Enum", "PAGE, PGA, MEDIA, SETTING, USER"],
        ["entityId", "String", "ID entitas yang diubah"],
        ["oldValue", "JSON (optional)", "Nilai sebelum perubahan"],
        ["newValue", "JSON (optional)", "Nilai sesudah perubahan"],
        ["ipAddress", "String (optional)", "Alamat IP pengguna"],
        ["userAgent", "String (optional)", "Browser/user agent"],
        ["createdAt", "DateTime", "Waktu aksi"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], audit_schema, [1.2, 1.8, 3.3])
    
    doc.add_paragraph()
    add_section_heading(doc, "Tabel: page_revisions", level=3)
    revisions_schema = [
        ["id", "String (UUID)", "Primary key"],
        ["pageId", "String (FK)", "Referensi ke halaman"],
        ["version", "Int", "Nomor versi"],
        ["content", "Text", "Snapshot konten pada versi ini"],
        ["changeSummary", "String (optional)", "Ringkasan perubahan"],
        ["createdById", "String (FK)", "User yang membuat revisi"],
        ["createdAt", "DateTime", "Waktu revisi"],
    ]
    create_styled_table(doc, ["Kolom", "Tipe", "Deskripsi"], revisions_schema, [1.2, 1.8, 3.3])
    
    # 3.4 API Endpoints
    add_section_heading(doc, "3.4 API Endpoints", level=2)
    
    add_body_paragraph(doc, "Berikut adalah daftar lengkap API endpoints yang tersedia dalam sistem:")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Authentication", bold=True, size=11, color=BURGUNDY)
    
    auth_endpoints = [
        ["POST", "/api/auth/login", "Login pengguna (email + password)"],
        ["POST", "/api/auth/logout", "Logout pengguna"],
        ["GET", "/api/auth/session", "Ambil data session aktif"],
        ["POST", "/api/auth/register", "Registrasi pengguna baru (Super Admin only)"],
    ]
    create_styled_table(doc, ["Method", "Endpoint", "Deskripsi"], auth_endpoints, [0.8, 2.2, 3.3])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Pages (Halaman)", bold=True, size=11, color=BURGUNDY)
    
    pages_endpoints = [
        ["GET", "/api/pages", "List semua halaman (paginated, filterable)"],
        ["GET", "/api/pages/:id", "Detail halaman berdasarkan ID"],
        ["GET", "/api/pages/slug/:slug", "Detail halaman berdasarkan slug"],
        ["POST", "/api/pages", "Buat halaman baru"],
        ["PUT", "/api/pages/:id", "Update halaman"],
        ["DELETE", "/api/pages/:id", "Hapus halaman"],
        ["POST", "/api/pages/:id/publish", "Publikasikan halaman"],
        ["POST", "/api/pages/:id/submit-review", "Kirim untuk review"],
        ["POST", "/api/pages/:id/approve", "Setujui halaman"],
        ["GET", "/api/pages/:id/revisions", "List revisi halaman"],
        ["GET", "/api/pages/:id/revisions/:version", "Detail revisi tertentu"],
    ]
    create_styled_table(doc, ["Method", "Endpoint", "Deskripsi"], pages_endpoints, [0.8, 2.7, 2.8])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "PGA (Pilar Fondasional)", bold=True, size=11, color=BURGUNDY)
    
    pga_endpoints = [
        ["GET", "/api/pgas", "List semua PGA (paginated, filterable by domain)"],
        ["GET", "/api/pgas/:id", "Detail PGA berdasarkan ID"],
        ["GET", "/api/pgas/code/:code", "Detail PGA berdasarkan kode (PGA-01)"],
        ["POST", "/api/pgas", "Buat PGA baru"],
        ["PUT", "/api/pgas/:id", "Update PGA"],
        ["DELETE", "/api/pgas/:id", "Hapus PGA"],
        ["GET", "/api/pgas/:id/cross-refs", "Referensi silang PGA"],
    ]
    create_styled_table(doc, ["Method", "Endpoint", "Deskripsi"], pga_endpoints, [0.8, 2.7, 2.8])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Domains", bold=True, size=11, color=BURGUNDY)
    
    domain_endpoints = [
        ["GET", "/api/domains", "List semua domain"],
        ["GET", "/api/domains/:id", "Detail domain"],
        ["GET", "/api/domains/:id/pgas", "List PGA dalam domain"],
        ["POST", "/api/domains", "Buat domain baru"],
        ["PUT", "/api/domains/:id", "Update domain"],
    ]
    create_styled_table(doc, ["Method", "Endpoint", "Deskripsi"], domain_endpoints, [0.8, 2.7, 2.8])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Media", bold=True, size=11, color=BURGUNDY)
    
    media_endpoints = [
        ["GET", "/api/media", "List semua media (paginated)"],
        ["POST", "/api/media/upload", "Upload file media"],
        ["GET", "/api/media/:id", "Detail media"],
        ["DELETE", "/api/media/:id", "Hapus media"],
    ]
    create_styled_table(doc, ["Method", "Endpoint", "Deskripsi"], media_endpoints, [0.8, 2.7, 2.8])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Users (Manajemen Pengguna)", bold=True, size=11, color=BURGUNDY)
    
    users_endpoints = [
        ["GET", "/api/users", "List semua pengguna (Super Admin only)"],
        ["GET", "/api/users/:id", "Detail pengguna"],
        ["POST", "/api/users", "Buat pengguna baru (Super Admin only)"],
        ["PUT", "/api/users/:id", "Update pengguna"],
        ["DELETE", "/api/users/:id", "Nonaktifkan pengguna"],
        ["PUT", "/api/users/:id/role", "Ubah role pengguna"],
    ]
    create_styled_table(doc, ["Method", "Endpoint", "Deskripsi"], users_endpoints, [0.8, 2.7, 2.8])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Settings & Audit", bold=True, size=11, color=BURGUNDY)
    
    settings_endpoints = [
        ["GET", "/api/settings", "List semua pengaturan"],
        ["PUT", "/api/settings/:key", "Update pengaturan"],
        ["GET", "/api/audit-log", "List audit log (paginated, filterable)"],
        ["GET", "/api/audit-log/:id", "Detail audit log"],
        ["GET", "/api/dashboard/stats", "Statistik dashboard"],
    ]
    create_styled_table(doc, ["Method", "Endpoint", "Deskripsi"], settings_endpoints, [0.8, 2.7, 2.8])
    
    # 3.5 Authentication
    add_section_heading(doc, "3.5 Authentication", level=2)
    
    add_body_paragraph(doc,
        "Sistem autentikasi menggunakan NextAuth.js v5 (Auth.js) dengan konfigurasi berikut:")
    
    auth_details = [
        ("Provider: ", "Credentials Provider (email + password)"),
        ("Session: ", "JWT-based session dengan refresh token rotation"),
        ("Password Hashing: ", "bcrypt dengan salt rounds = 12"),
        ("Session Duration: ", "24 jam (dapat di-configure via settings)"),
        ("CSRF Protection: ", "Built-in CSRF token dari NextAuth.js"),
        ("Rate Limiting: ", "Maksimal 5 percobaan login per 15 menit"),
        ("Secure Headers: ", "HTTP Strict Transport Security (HSTS), Content Security Policy"),
    ]
    for bold_part, normal_part in auth_details:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    doc.add_paragraph()
    add_body_paragraph(doc, "Role-Based Access Control (RBAC):")
    
    rbac_data = [
        ["Fitur", "Super Admin", "Editor", "Reviewer"],
        ["Dashboard", "✅ Lihat semua", "✅ Lihat miliknya", "✅ Lihat pending"],
        ["CRUD Halaman", "✅ Penuh", "✅ Create/Update", "❌ Read-only"],
        ["CRUD PGA", "✅ Penuh", "✅ Create/Update", "❌ Read-only"],
        ["Publish Konten", "✅ Langsung", "❌ Perlu review", "✅ Approve"],
        ["Manajemen User", "✅ Penuh", "❌", "❌"],
        ["Pengaturan Sistem", "✅ Penuh", "❌", "❌"],
        ["Audit Log", "✅ Lihat semua", "✅ Lihat miliknya", "✅ Lihat terkait"],
        ["Media Upload", "✅", "✅", "❌"],
    ]
    create_styled_table(doc, rbac_data[0], rbac_data[1:], [1.8, 1.2, 1.2, 1.2])
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 4: FITUR CMS (ADMIN PANEL)
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 4: FITUR CMS (ADMIN PANEL)", level=1)
    add_separator(doc)
    
    # 4.1 Dashboard
    add_section_heading(doc, "4.1 Dashboard", level=2)
    
    add_body_paragraph(doc,
        "Dashboard adalah halaman utama yang ditampilkan setelah admin berhasil login. Dashboard "
        "menyediakan gambaran umum status seluruh konten dalam sistem.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Komponen Dashboard:", bold=True, size=11, color=BURGUNDY)
    
    dashboard_items = [
        ("Statistik Ringkasan: ", "Total halaman, PGA yang dipublikasikan, draft pending, media files count."),
        ("Kartu Status Konten: ", "4 kartu yang menampilkan jumlah konten berdasarkan status (Draft, Pending Review, Approved, Published)."),
        ("Aktivitas Terbaru: ", "Timeline 10 perubahan terakhir dengan informasi user, aksi, dan timestamp."),
        ("Quick Actions: ", "Tombol cepat untuk aksi paling umum: Tambah Halaman, Tambah PGA, Upload Media."),
        ("Grafik Domain: ", "Visualisasi jumlah PGA per domain dalam bentuk bar chart."),
        ("Notifikasi: ", "Badge counter untuk konten yang menunggu review."),
    ]
    for bold_part, normal_part in dashboard_items:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 4.2 Manajemen Halaman
    add_section_heading(doc, "4.2 Manajemen Halaman", level=2)
    
    add_body_paragraph(doc,
        "Fitur ini memungkinkan admin untuk mengelola seluruh halaman flipbook. Setiap halaman memiliki "
        "metadata, konten, dan status yang dapat dikelola secara individual.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Fitur Manajemen Halaman:", bold=True, size=11, color=BURGUNDY)
    
    page_mgmt_items = [
        ("List View: ", "Tabel daftar semua halaman dengan kolom: No., Judul, Tipe, Status, Domain, Penulis, Terakhir Diupdate."),
        ("Filter & Search: ", "Pencarian berdasarkan judul, filter berdasarkan tipe, status, dan domain."),
        ("Sort: ", "Urutkan berdasarkan nomor halaman, judul, status, atau tanggal update."),
        ("Bulk Actions: ", "Aksi massal untuk publish, unpublish, atau hapus beberapa halaman sekaligus."),
        ("Page Navigation: ", "Pagination dengan opsi 10/25/50 item per halaman."),
        ("Status Indicators: ", "Badge berwarna untuk setiap status: Draft (abu), Pending (kuning), Approved (hijau), Published (biru)."),
    ]
    for bold_part, normal_part in page_mgmt_items:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 4.3 Manajemen PGA-72
    add_section_heading(doc, "4.3 Manajemen PGA-72", level=2)
    
    add_body_paragraph(doc,
        "Editor khusus untuk mengelola 72 Pilar Fondasional. Setiap PGA memiliki struktur data yang "
        "lebih kompleks dibanding halaman biasa, termasuk dimensi, prinsip, dan referensi silang.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Form Editor PGA meliputi field:", bold=True, size=11, color=BURGUNDY)
    
    pga_fields = [
        "Kode PGA (auto-generated: PGA-01 s/d PGA-72)",
        "Nama PGA (Bahasa Indonesia) — wajib diisi",
        "Nama PGA (Bahasa Inggris) — opsional",
        "Domain PGA (dropdown D1-D9) — wajib dipilih",
        "Vision Statement — textarea untuk pernyataan visi PGA",
        "Deskripsi — rich text editor untuk konten lengkap",
        "Dimensi — dynamic list, bisa tambah/hapus dimensi",
        "Prinsip — dynamic list, bisa tambah/hapus prinsip",
        "Referensi Silang — multi-select dari daftar PGA lain",
        "Warna Representasi — color picker (default dari domain)",
    ]
    for item in pga_fields:
        add_bullet_point(doc, item)
    
    # 4.4 Editor Konten
    add_section_heading(doc, "4.4 Editor Konten", level=2)
    
    add_body_paragraph(doc,
        "Editor konten adalah komponen inti dari CMS yang memungkinkan pengguna untuk menulis dan "
        "mengedit konten secara visual tanpa perlu memahami HTML atau Markdown.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Fitur Editor Konten:", bold=True, size=11, color=BURGUNDY)
    
    editor_features = [
        ("Rich Text Formatting: ", "Bold, italic, underline, strikethrough, headings, lists, blockquote, code blocks."),
        ("MDX Support: ", "Konten disimpan dalam format MDX yang mendukung komponen React khusus."),
        ("Image Embedding: ", "Insert gambar dari media library atau upload langsung ke editor."),
        ("Table Support: ", "Membuat dan mengedit tabel langsung di editor."),
        ("Preview Mode: ", "Preview konten dalam tampilan flipbook tanpa meninggalkan editor."),
        ("Auto-save: ", "Konten otomatis tersimpan sebagai draft setiap 30 detik."),
        ("Version History: ", "Menampilkan daftar versi sebelumnya dengan kemampuan restore."),
        ("Undo/Redo: ", "Undo dan redo tanpa batas selama sesi editing."),
        ("Keyboard Shortcuts: ", "Shortcut keyboard untuk formatting cepat (Ctrl+B, Ctrl+I, dll)."),
        ("Fullscreen Mode: ", "Mode layar penuh untuk fokus pada penulisan."),
    ]
    for bold_part, normal_part in editor_features:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 4.5 Media Manager
    add_section_heading(doc, "4.5 Media Manager", level=2)
    
    add_body_paragraph(doc,
        "Media Manager mengelola semua file media (gambar, dokumen) yang digunakan dalam konten flipbook.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Fitur Media Manager:", bold=True, size=11, color=BURGUNDY)
    
    media_features = [
        ("Upload: ", "Drag & drop atau click to upload, mendukung JPEG, PNG, SVG, WebP, PDF."),
        ("Size Limits: ", "Maksimal 10MB per file, 100MB total storage (configurable)."),
        ("Image Optimization: ", "Otomatis resize dan compress gambar saat upload."),
        ("Grid/List View: ", "Tampilan grid thumbnail atau list detail."),
        ("Search & Filter: ", "Pencarian berdasarkan nama file, filter berdasarkan tipe dan tanggal."),
        ("Folder Organization: ", "Pengelompokan file dalam folder (opsional)."),
        ("Copy URL: ", "Salin URL file untuk digunakan di konten."),
        ("Delete: ", "Hapus file dengan konfirmasi (cek penggunaan di konten sebelum hapus)."),
    ]
    for bold_part, normal_part in media_features:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 4.6 Theme & Design
    add_section_heading(doc, "4.6 Theme & Design", level=2)
    
    add_body_paragraph(doc,
        "Pengaturan tema dan desain memungkinkan admin untuk menyesuaikan tampilan visual flipbook "
        "tanpa mengubah kode. Semua pengaturan disimpan di database dan diterapkan secara dinamis.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Pengaturan Tema:", bold=True, size=11, color=BURGUNDY)
    
    theme_features = [
        ("Color Scheme: ", "Color picker untuk warna primer (burgundy #5E2129), sekunder (gold #C5A059), aksen, dan background (parchment #F5F1E8)."),
        ("Typography: ", "Pemilihan font untuk heading, body text, dan UI dari library font yang tersedia."),
        ("Batik Pattern: ", "Toggle on/off dan pilihan pola batik Kawung untuk elemen dekoratif."),
        ("Page Transitions: ", "Pilihan animasi transisi halaman flipbook."),
        ("Sound Effects: ", "Toggle on/off efek suara saat membalik halaman."),
        ("Logo & Favicon: ", "Upload logo KNMP dan favicon untuk browser tab."),
        ("Custom CSS: ", "Textarea untuk CSS kustom (hanya Super Admin)."),
    ]
    for bold_part, normal_part in theme_features:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 4.7 Pengaturan Situs
    add_section_heading(doc, "4.7 Pengaturan Situs", level=2)
    
    add_body_paragraph(doc,
        "Pengaturan situs mencakup konfigurasi metadata dan SEO yang mempengaruhi tampilan dan "
        "discoverability flipbook.")
    
    site_settings = [
        ["site_title", "Judul Situs", "KNBMP PGA-72 Flipbook"],
        ["site_description", "Deskripsi Situs", "Dokumen Pilar Fondasional KNBMP"],
        ["site_url", "URL Situs", "https://flipbook.knmp.or.id"],
        ["og_image", "Open Graph Image", "Path ke gambar OG"],
        ["favicon", "Favicon", "Path ke favicon"],
        ["analytics_id", "Google Analytics ID", "UA-XXXXXX-X"],
        ["show_page_numbers", "Tampilkan Nomor Halaman", "true/false"],
        ["enable_sound", "Efek Suara Flipbook", "true/false"],
        ["auto_play_animation", "Auto-play Animasi Cover", "true/false"],
    ]
    create_styled_table(doc, ["Key", "Label", "Default Value"], site_settings, [1.5, 2.0, 2.8])
    
    # 4.8 Audit Log
    add_section_heading(doc, "4.8 Audit Log", level=2)
    
    add_body_paragraph(doc,
        "Audit log mencatat setiap aktivitas yang terjadi dalam sistem untuk keperluan akuntabilitas "
        "dan troubleshooting. Semua entri bersifat read-only dan tidak dapat dihapus.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Informasi yang Tercatat:", bold=True, size=11, color=BURGUNDY)
    
    audit_features = [
        "Timestamp (waktu aksi dalam UTC+7)",
        "User yang melakukan aksi (nama dan email)",
        "Tipe aksi (CREATE, UPDATE, DELETE, PUBLISH, LOGIN, LOGOUT)",
        "Tipe entitas yang diubah (PAGE, PGA, MEDIA, SETTING, USER)",
        "ID entitas yang diubah",
        "Nilai sebelum perubahan (diff/old value)",
        "Nilai sesudah perubahan (new value)",
        "Alamat IP pengguna",
        "User Agent (browser info)",
    ]
    for item in audit_features:
        add_bullet_point(doc, item)
    
    add_body_paragraph(doc,
        "Audit log dapat difilter berdasarkan user, tipe aksi, tipe entitas, dan rentang tanggal. "
        "Export ke CSV tersedia untuk analisis lebih lanjut.")
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 5: STRUKTUR KONTEN
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 5: STRUKTUR KONTEN", level=1)
    add_separator(doc)
    
    # 5.1 Peta Halaman
    add_section_heading(doc, "5.1 Peta Halaman", level=2)
    
    add_body_paragraph(doc,
        "Flipbook KNBMP PGA-72 terdiri dari total 96 halaman dengan struktur sebagai berikut. "
        "Setiap halaman memiliki tipe, nomor, dan metadata yang tersimpan di database CMS.")
    
    page_map = [
        ["1", "COVER", "Halaman Sampul", "—", "Desain heritage dengan batik Kawung"],
        ["2", "FOREWORD", "Kata Pengantar KP-1", "—", "Pendahuluan oleh Ketua Umum"],
        ["3", "FOREWORD", "Kata Pengantar KP-2", "—", "Pendahuluan oleh Sekretaris Jenderal"],
        ["4", "FOREWORD", "Kata Pengantar KP-3", "—", "Pendahuluan oleh Bendahara Umum"],
        ["5", "FOREWORD", "Kata Pengantar KP-4", "—", "Pendahuluan oleh Dewan Pengawas"],
        ["6", "FOREWORD", "Kata Pengantar KP-5", "—", "Kata Pengantar tambahan"],
        ["7", "FOREWORD", "Kata Pengantar KP-6", "—", "Kata Pengantar tambahan"],
        ["8", "FOREWORD", "Kata Pengantar KP-7", "—", "Kata Pengantar tambahan"],
        ["9", "MUKADIMAH", "Mukadimah Bagian 1", "—", "Pendahuluan Mukadimah"],
        ["10", "MUKADIMAH", "Mukadimah Bagian 2", "—", "Latar Belakang"],
        ["11", "MUKADIMAH", "Mukadimah Bagian 3", "—", "Tujuan & Sasaran"],
        ["12", "MUKADIMAH", "Mukadimah Bagian 4", "—", "Ruang Lingkup"],
        ["13", "TOC", "Daftar Isi - Domain 1-3", "D1-D3", "Identity, Strategy, Governance"],
        ["14", "TOC", "Daftar Isi - Domain 4-6", "D4-D6", "Product, Operations, People"],
        ["15", "TOC", "Daftar Isi - Domain 7-9", "D7-D9", "Finance, Risk, Innovation"],
        ["16-21", "TOC", "Daftar Isi Detail", "D1-D9", "Detail setiap domain dan PGA"],
    ]
    create_styled_table(doc, ["Hal.", "Tipe", "Judul", "Domain", "Catatan"], page_map, [0.6, 0.9, 2.0, 0.8, 2.0])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Halaman PGA (22-93): 72 Pilar Fondasional", bold=True, size=11, color=BURGUNDY)
    add_body_paragraph(doc,
        "Halaman 22 hingga 93 berisi 72 PGA yang dikelompokkan ke dalam 9 domain. Detail domain "
        "dan PGA disajikan pada sub-bab berikutnya.")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Halaman Penutup:", bold=True, size=11, color=BURGUNDY)
    
    closing_pages = [
        ["94", "PHILOSOPHY", "Filsafat KNBMP", "—", "Pernyataan filosofis gerakan"],
        ["95", "PACT", "Pakta Integritas", "—", "Pakta integritas anggota dan pengurus"],
        ["96", "BACK_COVER", "Halaman Belakang", "—", "Informasi kontak dan QR code"],
    ]
    create_styled_table(doc, ["Hal.", "Tipe", "Judul", "Domain", "Catatan"], closing_pages, [0.6, 0.9, 2.0, 0.8, 2.0])
    
    # 5.2 9 Domain Detail
    add_section_heading(doc, "5.2 9 Domain Detail", level=2)
    
    add_body_paragraph(doc,
        "72 Pilar Fondasional (PGA) dikelompokkan ke dalam 9 domain fungsional. Setiap domain "
        "memiliki 8 PGA dengan kode warna yang representatif:")
    
    domain_details = [
        ["D1", "Identity & Civilization", "Identitas & Peradaban", "#C4952A", "PGA-01 s/d PGA-08", "Fondasi identitas organisasi dan peradaban koperasi"],
        ["D2", "Strategy & Direction", "Strategi & Arah", "#1565C0", "PGA-09 s/d PGA-16", "Arah strategis dan rencana jangka panjang"],
        ["D3", "Governance & Legal", "Tata Kelola & Hukum", "#6A1B9A", "PGA-17 s/d PGA-24", "Kerangka tata kelola dan kepatuhan hukum"],
        ["D4", "Product/Service/Solution", "Produk/Layanan/Solusi", "#008F3D", "PGA-25 s/d PGA-32", "Portofolio produk dan layanan koperasi"],
        ["D5", "Operations & Infrastructure", "Operasi & Infrastruktur", "#E65100", "PGA-33 s/d PGA-40", "Infrastruktur operasional dan teknologi"],
        ["D6", "People & Community", "Sumber Daya Manusia & Komunitas", "#AD1457", "PGA-41 s/d PGA-48", "Pengembangan SDM dan penguatan komunitas"],
        ["D7", "Finance & Sustainability", "Keuangan & Keberlanjutan", "#00838F", "PGA-49 s/d PGA-56", "Manajemen keuangan dan keberlanjutan"],
        ["D8", "Risk & Resilience", "Risiko & Ketahanan", "#4527A0", "PGA-57 s/d PGA-64", "Manajemen risiko dan ketahanan organisasi"],
        ["D9", "Innovation & Future", "Inovasi & Masa Depan", "#BF360C", "PGA-65 s/d PGA-72", "Inovasi dan visi masa depan koperasi"],
    ]
    create_styled_table(doc, ["Kode", "Nama (EN)", "Nama (ID)", "Warna", "PGA", "Deskripsi"], domain_details, 
                       [0.5, 1.2, 1.2, 0.7, 1.0, 1.7])
    
    # 5.3 Struktur Data Per PGA
    add_section_heading(doc, "5.3 Struktur Data Per PGA", level=2)
    
    add_body_paragraph(doc,
        "Setiap PGA memiliki struktur data yang terstandarisasi. Berikut adalah contoh struktur "
        "data lengkap untuk satu PGA:")
    
    add_info_box(doc, "CONTOH: PGA-01 (Identitas Koperasi)",
        """{
  "id": "uuid-pga-01",
  "code": "PGA-01",
  "name": "Identitas Koperasi",
  "nameEn": "Cooperative Identity",
  "description": "Fondasi identitas yang membedakan koperasi dari entitas bisnis lain...",
  "domainId": "domain-d1",
  "domainCode": "D1",
  "vision": "Menjadi koperasi dengan identitas yang kuat dan diakui secara nasional",
  "orderInDomain": 1,
  "color": "#C4952A",
  "dimensions": [
    "Nilai-nilai Dasar Koperasi",
    "Prinsip-prinsip Koperasi",
    "Identitas Visual & Branding"
  ],
  "principles": [
    "Keanggotaan terbuka dan sukarela",
    "Pengelolaan demokratis oleh anggota",
    "Partisipasi ekonomi anggota"
  ],
  "crossReferences": ["PGA-09", "PGA-17", "PGA-41"]
}""")
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 6: DESAIN ANTARMUKA
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 6: DESAIN ANTARMUKA", level=1)
    add_separator(doc)
    
    # 6.1 Design System
    add_section_heading(doc, "6.1 Design System", level=2)
    
    add_body_paragraph(doc,
        "Design system KNBMP PGA-72 mengusung tema heritage yang terinspirasi dari budaya "
        "Indonesia, khususnya motif batik dan warna-warna tradisional. Tema ini mencerminkan "
        "identitas, kedaulatan, dan kekayaan budaya Nusantara.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Palet Warna Utama:", bold=True, size=11, color=BURGUNDY)
    
    color_palette = [
        ["Burgundy (Primer)", "#5E2129", "Warna utama untuk heading, tombol, dan elemen penting"],
        ["Gold (Sekunder)", "#C5A059", "Aksen dekoratif, border, dan elemen highlight"],
        ["Parchment (Background)", "#F5F1E8", "Warna latar belakang utama flipbook"],
        ["Dark Text", "#2D2D2D", "Warna teks utama untuk keterbacaan"],
        ["Medium Text", "#555555", "Warna teks sekunder untuk caption dan deskripsi"],
        ["Light Gold (Hover)", "#F5EED5", "Background hover untuk elemen interaktif"],
    ]
    create_styled_table(doc, ["Nama Warna", "Hex Code", "Penggunaan"], color_palette, [1.5, 1.0, 3.8])
    
    # 6.2 Typography
    add_section_heading(doc, "6.2 Typography", level=2)
    
    add_body_paragraph(doc,
        "Sistem tipografi menggunakan kombinasi font serif dan sans-serif untuk menciptakan "
        "kesan mewah namun tetap modern dan mudah dibaca.")
    
    typo_data = [
        ["DM Serif Display", "Heading Utama", "Judul bab, judul halaman, elemen hero", "36-48px"],
        ["Libre Baskerville", "Body Text", "Paragraf utama, deskripsi panjang", "14-16px"],
        ["Inter", "UI Elements", "Label, tombol, navigasi, form input", "12-14px"],
        ["Courier New", "Vision Statement", "Kutipan, pernyataan visi PGA", "14-16px"],
        ["DM Serif Display", "Cover Title", "Judul di halaman sampul", "42-56px"],
    ]
    create_styled_table(doc, ["Font", "Kategori", "Penggunaan", "Ukuran"], typo_data, [1.3, 1.1, 2.8, 0.9])
    
    # 6.3 Batik Kawung Pattern
    add_section_heading(doc, "6.3 Batik Kawung Pattern", level=2)
    
    add_body_paragraph(doc,
        "Motif Batik Kawung dipilih sebagai elemen dekoratif utama karena memiliki makna filosofis "
        "yang mendalam dalam budaya Jawa. Motif Kawung melambangkan harapan agar pemimpin menjadi "
        "pemimpin yang adil kepada rakyatnya — nilai yang sangat relevan dengan prinsip koperasi.")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Penerapan Batik Kawung:", bold=True, size=11, color=BURGUNDY)
    
    batik_usage = [
        ("Halaman Sampul: ", "Background pattern dengan opacity rendah sebagai dekorasi."),
        ("Header Setiap Bab: ", "Strip batik horizontal sebagai pembatas antar bab."),
        ("Border Halaman: ", "Frame batik tipis di sekeliling konten halaman."),
        ("Divider Section: ", "Motif batik sebagai pemisah antar section dalam halaman."),
        ("Watermark: ", "Pattern batik semi-transparan sebagai watermark halaman."),
    ]
    for bold_part, normal_part in batik_usage:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 6.4 Responsive Design
    add_section_heading(doc, "6.4 Responsive Design", level=2)
    
    add_body_paragraph(doc,
        "Flipbook dan Admin Panel dirancang dengan pendekatan mobile-first untuk memastikan "
        "aksesibilitas di berbagai ukuran layar.")
    
    responsive_breakpoints = [
        ["Mobile S", "320px - 374px", "1 kolom, flipbook single page, font 14px"],
        ["Mobile L", "375px - 424px", "1 kolom, flipbook single page, font 14px"],
        ["Tablet", "425px - 768px", "2 kolom partial, flipbook single page, font 15px"],
        ["Laptop", "769px - 1024px", "2 kolom, flipbook spread ready, font 16px"],
        ["Desktop", "1025px - 1439px", "Full layout, flipbook double spread, font 16px"],
        ["Desktop XL", "1440px+", "Full layout with margins, flipbook centered, font 16px"],
    ]
    create_styled_table(doc, ["Breakpoint", "Lebar", "Layout"], responsive_breakpoints, [1.2, 1.3, 3.8])
    
    # 6.5 Flipbook Interactions
    add_section_heading(doc, "6.5 Flipbook Interactions", level=2)
    
    add_body_paragraph(doc,
        "Flipbook dirancang untuk memberikan pengalaman membaca digital yang mendekati "
        "pengalaman membaca buku fisik, dengan sentuhan modern.")
    
    interactions = [
        ("3D Page Flip: ", "Animasi page flip 3D yang realistis menggunakan CSS transforms dan perspective."),
        ("Sound Effects: ", "Efek suara halaman kertas yang dapat di-toggle on/off."),
        ("Swipe Gesture: ", "Geser layar untuk berpindah halaman pada perangkat touch."),
        ("Keyboard Navigation: ", "Arrow keys untuk berpindah halaman, Home/End untuk lompat ke awal/akhir."),
        ("Page Thumbnail Strip: ", "Strip thumbnail di bawah flipbook untuk navigasi cepat."),
        ("Zoom: ", "Pinch-to-zoom pada mobile dan scroll zoom pada desktop."),
        ("Table of Contents Sidebar: ", "Sidebar navigasi yang dapat dibuka/tutup untuk akses cepat."),
        ("Search: ", "Pencarian teks dalam seluruh konten flipbook."),
        ("Bookmark: ", "Penanda halaman untuk akses cepat ke halaman favorit."),
        ("Full-screen Mode: ", "Mode layar penuh untuk pengalaman membaca tanpa gangguan."),
    ]
    for bold_part, normal_part in interactions:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 7: IMPLEMENTASI
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 7: IMPLEMENTASI", level=1)
    add_separator(doc)
    
    # 7.1 Roadmap 3 Fase
    add_section_heading(doc, "7.1 Roadmap 3 Fase", level=2)
    
    add_body_paragraph(doc,
        "Implementasi CMS KNBMP PGA-72 dibagi menjadi 3 fase dengan total durasi 6 minggu. "
        "Setiap fase memiliki deliverable yang terukur dan dapat di-demo-kan.")
    
    # Fase 1
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=312)
    add_formatted_run(p, "FASE 1: Foundation (Minggu 1-2)", bold=True, size=12, color=BURGUNDY)
    
    add_body_paragraph(doc, "Fokus: Database, API, dan Autentikasi")
    
    fase1_tasks = [
        ["1.1", "Setup proyek Next.js 16 dengan TypeScript", "Hari 1", "High"],
        ["1.2", "Desain dan implementasi database schema (Prisma)", "Hari 1-2", "High"],
        ["1.3", "Implementasi API endpoints untuk Pages (CRUD)", "Hari 2-4", "High"],
        ["1.4", "Implementasi API endpoints untuk PGA (CRUD)", "Hari 3-5", "High"],
        ["1.5", "Implementasi API endpoints untuk Domains", "Hari 4", "High"],
        ["1.6", "Setup NextAuth.js dengan credentials provider", "Hari 5-6", "High"],
        ["1.7", "Implementasi RBAC middleware", "Hari 6-7", "High"],
        ["1.8", "Seed data: 9 domains + 72 PGA + sample pages", "Hari 7-8", "Medium"],
        ["1.9", "API endpoint untuk Media (upload/CRUD)", "Hari 8-9", "Medium"],
        ["1.10", "Unit tests untuk API endpoints", "Hari 9-10", "Medium"],
    ]
    create_styled_table(doc, ["ID", "Task", "Timeline", "Prioritas"], fase1_tasks, [0.5, 3.3, 1.0, 0.8])
    
    doc.add_paragraph()
    # Fase 2
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=312)
    add_formatted_run(p, "FASE 2: Admin Panel (Minggu 3-4)", bold=True, size=12, color=BURGUNDY)
    
    add_body_paragraph(doc, "Fokus: Antarmuka Manajemen Konten")
    
    fase2_tasks = [
        ["2.1", "Setup layout Admin Panel (sidebar, header, breadcrumbs)", "Hari 11-12", "High"],
        ["2.2", "Halaman Login dan Register", "Hari 11", "High"],
        ["2.3", "Dashboard dengan statistik dan quick actions", "Hari 12-13", "High"],
        ["2.4", "Manajemen Halaman (list, filter, search)", "Hari 13-15", "High"],
        ["2.5", "Form Editor Halaman (create/edit)", "Hari 14-16", "High"],
        ["2.6", "Editor Konten (rich text, MDX, preview)", "Hari 15-18", "High"],
        ["2.7", "Manajemen PGA-72 (list, form editor)", "Hari 16-18", "High"],
        ["2.8", "Media Manager (upload, grid, delete)", "Hari 18-19", "Medium"],
        ["2.9", "Theme & Design settings panel", "Hari 19-20", "Medium"],
        ["2.10", "Audit Log viewer dengan filter", "Hari 20", "Medium"],
    ]
    create_styled_table(doc, ["ID", "Task", "Timeline", "Prioritas"], fase2_tasks, [0.5, 3.3, 1.0, 0.8])
    
    doc.add_paragraph()
    # Fase 3
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=312)
    add_formatted_run(p, "FASE 3: Integration & Testing (Minggu 5-6)", bold=True, size=12, color=BURGUNDY)
    
    add_body_paragraph(doc, "Fokus: Integrasi Flipbook, Testing, dan Deployment")
    
    fase3_tasks = [
        ["3.1", "Integrasi CMS data dengan Flipbook frontend", "Hari 21-23", "High"],
        ["3.2", "Implementasi auto-update (ISR / revalidation)", "Hari 22-23", "High"],
        ["3.3", "Workflow: Draft → Review → Publish", "Hari 23-24", "High"],
        ["3.4", "Version history dan restore functionality", "Hari 24", "Medium"],
        ["3.5", "End-to-end testing seluruh flow", "Hari 25-26", "High"],
        ["3.6", "Performance optimization dan caching", "Hari 26-27", "Medium"],
        ["3.7", "Security audit dan penetration testing", "Hari 27", "High"],
        ["3.8", "User acceptance testing (UAT)", "Hari 28", "High"],
        ["3.9", "Documentation dan user guide", "Hari 28-29", "Medium"],
        ["3.10", "Deployment ke production (Vercel)", "Hari 29-30", "High"],
    ]
    create_styled_table(doc, ["ID", "Task", "Timeline", "Prioritas"], fase3_tasks, [0.5, 3.3, 1.0, 0.8])
    
    # 7.2 Deployment
    add_section_heading(doc, "7.2 Deployment", level=2)
    
    add_body_paragraph(doc,
        "Sistem CMS KNBMP PGA-72 dirancang untuk dapat di-deploy ke dua opsi utama:")
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Opsi 1: Vercel (Recommended)", bold=True, size=11, color=BURGUNDY)
    
    vercel_items = [
        "Automatic deployments dari Git repository",
        "Edge functions untuk API routes",
        "Built-in CDN untuk static assets",
        "Free tier tersedia untuk prototyping",
        "Automatic HTTPS dan SSL certificates",
    ]
    for item in vercel_items:
        add_bullet_point(doc, item)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Opsi 2: Self-Hosted (Standalone)", bold=True, size=11, color=BURGUNDY)
    
    standalone_items = [
        "Docker containerization untuk portabilitas",
        "Nginx sebagai reverse proxy",
        "PM2 untuk process management",
        "SQLite database file untuk kemudahan backup",
        "Cron job untuk automated backup harian",
    ]
    for item in standalone_items:
        add_bullet_point(doc, item)
    
    # 7.3 Backup & Recovery
    add_section_heading(doc, "7.3 Backup & Recovery", level=2)
    
    add_body_paragraph(doc,
        "Strategi backup dirancang untuk memastikan tidak ada data yang hilang dalam kondisi apapun.")
    
    backup_data = [
        ["Database (SQLite)", "Setiap 6 jam", "7 hari (rolling)", "rsync ke cloud storage"],
        ["Media Files", "Setiap hari (00:00)", "30 hari (rolling)", "rsync ke cloud storage"],
        ["Konfigurasi Sistem", "Setiap perubahan", "Versi di Git", "Git commit otomatis"],
        ["Audit Log", "Real-time", "1 tahun", "Separate SQLite file"],
        ["Full Backup", "Mingguan (Minggu)", "4 minggu (rolling)", "Compressed archive"],
    ]
    create_styled_table(doc, ["Data", "Frekuensi", "Retensi", "Metode"], backup_data, [1.3, 1.3, 1.3, 2.4])
    
    # 7.4 Monitoring
    add_section_heading(doc, "7.4 Monitoring", level=2)
    
    add_body_paragraph(doc,
        "Sistem monitoring memastikan ketersediaan dan performa CMS tetap optimal.")
    
    monitoring_items = [
        ("Error Tracking: ", "Sentry atau equivalent untuk menangkap dan melaporkan error secara real-time."),
        ("Performance: ", "Lighthouse CI untuk monitoring Core Web Vitals (LCP, FID, CLS)."),
        ("Uptime: ", "Health check endpoint (/api/health) dengan monitoring interval 1 menit."),
        ("Resource Usage: ", "CPU, memory, dan disk usage monitoring via PM2 atau Docker stats."),
        ("API Response Time: ", "Logging response time untuk setiap API endpoint."),
        ("Database Size: ", "Monitoring ukuran file SQLite untuk mengantisipasi growth."),
        ("Security Alerts: ", "Notifikasi untuk login attempts yang gagal berulang kali."),
    ]
    for bold_part, normal_part in monitoring_items:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 8: KEAMANAN
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 8: KEAMANAN", level=1)
    add_separator(doc)
    
    add_body_paragraph(doc,
        "Keamanan adalah prioritas tertinggi dalam pengembangan CMS KNBMP PGA-72, mengingat "
        "data yang dikelola bersifat strategis dan berdampak pada 270 juta rakyat Indonesia.")
    
    # 8.1 Autentikasi
    add_section_heading(doc, "8.1 Autentikasi", level=2)
    
    auth_security = [
        ("NextAuth.js v5: ", "Menggunakan Auth.js terbaru dengan security best practices."),
        ("JWT Sessions: ", "Token JWT yang di-encode dan di-signed menggunakan secret key yang kuat."),
        ("Bcrypt Hashing: ", "Password di-hash menggunakan bcrypt dengan cost factor 12."),
        ("Secure Cookies: ", "Session cookies menggunakan flag HttpOnly, Secure, dan SameSite=Strict."),
        ("Token Expiration: ", "Akses token berlaku 24 jam, refresh token 7 hari."),
        ("Password Policy: ", "Minimum 8 karakter, kombinasi huruf besar, kecil, angka, dan simbol."),
    ]
    for bold_part, normal_part in auth_security:
        add_bullet_point(doc, normal_part, bold_prefix=bold_part)
    
    # 8.2 Otorisasi
    add_section_heading(doc, "8.2 Otorisasi", level=2)
    
    add_body_paragraph(doc,
        "Sistem otorisasi menggunakan Role-Based Access Control (RBAC) dengan 3 level peran:")
    
    rbac_details = [
        ["SUPER_ADMIN", "Akses penuh ke seluruh fitur sistem", "Hanya 1-2 orang"],
        ["EDITOR", "CRUD konten, upload media, lihat audit log miliknya", "Tim konten (3-5 orang)"],
        ["REVIEWER", "Review dan approve konten, lihat audit log terkait", "Pengawas (2-3 orang)"],
    ]
    create_styled_table(doc, ["Role", "Akses", "Kuota"], rbac_details, [1.2, 3.3, 1.8])
    
    add_body_paragraph(doc,
        "Implementasi middleware Next.js yang memeriksa role pada setiap request ke protected routes. "
        "Unauthorized access akan di-redirect ke halaman login dengan pesan error yang informatif.")
    
    # 8.3 Validasi Input
    add_section_heading(doc, "8.3 Validasi Input", level=2)
    
    add_body_paragraph(doc,
        "Semua input dari pengguna divalidasi menggunakan Zod schema sebelum diproses. "
        "Validasi dilakukan di kedua sisi: client-side (untuk UX) dan server-side (untuk keamanan).")
    
    validation_examples = [
        "Email: valid email format (z.string().email())",
        "Password: min 8 chars, 1 uppercase, 1 lowercase, 1 number (z.string().min(8).regex())",
        "Page Title: min 1 char, max 200 chars",
        "Page Content: max 100KB per halaman",
        "PGA Code: format PGA-XX (z.string().regex(/^PGA-\\d{2}$/))",
        "Color: valid hex color (z.string().regex(/^#[0-9A-Fa-f]{6}$/))",
        "File Upload: max 10MB, allowed MIME types only",
        "SQL Injection: Prisma ORM secara otomatis mencegah SQL injection",
        "XSS: Semua output di-escape secara default oleh React",
    ]
    for item in validation_examples:
        add_bullet_point(doc, item)
    
    # 8.4 Rate Limiting
    add_section_heading(doc, "8.4 Rate Limiting", level=2)
    
    rate_limits = [
        ["POST /api/auth/login", "5 requests / 15 menit", "Per IP address"],
        ["POST /api/auth/register", "3 requests / 60 menit", "Per IP address"],
        ["POST /api/pages", "30 requests / 15 menit", "Per user session"],
        ["PUT /api/pages/:id", "60 requests / 15 menit", "Per user session"],
        ["POST /api/media/upload", "10 requests / 15 menit", "Per user session"],
        ["GET /api/*", "100 requests / 15 menit", "Per IP address"],
    ]
    create_styled_table(doc, ["Endpoint", "Limit", "Basis"], rate_limits, [2.0, 2.0, 2.3])
    
    # 8.5 Audit Trail
    add_section_heading(doc, "8.5 Audit Trail", level=2)
    
    add_body_paragraph(doc,
        "Setiap aksi yang dilakukan dalam CMS dicatat dalam audit log untuk keperluan:")
    
    audit_purposes = [
        "Akuntabilitas: Melacak siapa yang melakukan perubahan dan kapan",
        "Forensik: Investigasi insiden keamanan",
        "Compliance: Memenuhi kebutuhan audit organisasi",
        "Recovery: Memulihkan data ke kondisi sebelumnya",
    ]
    for item in audit_purposes:
        add_bullet_point(doc, item)
    
    add_body_paragraph(doc,
        "Audit log bersifat immutable (tidak dapat diubah atau dihapus) dan disimpan dalam "
        "tabel terpisah. Export audit log tersedia dalam format CSV dan JSON.")
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 9: RISIKO & MITIGASI
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 9: RISIKO & MITIGASI", level=1)
    add_separator(doc)
    
    add_body_paragraph(doc,
        "Berikut adalah identifikasi risiko dan strategi mitigasi untuk proyek CMS KNBMP PGA-72:")
    
    risk_data = [
        ["Data Loss", "Kritis", "Sedang", "Automated backup 6 jam sekali, snapshot mingguan, versioning konten"],
        ["Unauthorized Access", "Kritis", "Rendah", "RBAC, JWT sessions, bcrypt, secure cookies, rate limiting"],
        ["SQL Injection", "Tinggi", "Rendah", "Prisma ORM parameterized queries, Zod validation"],
        ["XSS Attack", "Tinggi", "Rendah", "React auto-escaping, Content Security Policy headers"],
        ["Performance Degradation", "Sedang", "Sedang", "ISR, caching, CDN, database indexing, query optimization"],
        ["Single Point of Failure", "Sedang", "Sedang", "Health monitoring, automated restart (PM2), redundant backup"],
        ["Scope Creep", "Sedang", "Tinggi", "Strict PRD adherence, change request process, phase-based delivery"],
        ["Developer Dependency", "Tinggi", "Tinggi", "CMS enables non-technical users, documentation, user guide"],
        ["Budget Overrun", "Sedang", "Sedang", "Fixed scope per phase, MVP-first approach, open-source tools"],
        ["Technology Obsolescence", "Rendah", "Rendah", "Using mature tech (Next.js, React, Prisma), regular updates"],
        ["User Adoption", "Sedang", "Sedang", "Intuitive UI, training sessions, video tutorials, user guide"],
        ["Content Quality", "Sedang", "Sedang", "Draft → Review → Publish workflow, version control"],
        ["Data Migration Issues", "Sedang", "Rendah", "Clean schema design, migration scripts, rollback plan"],
        ["Third-party Service Outage", "Rendah", "Rendah", "Self-contained SQLite, minimal external dependencies"],
    ]
    create_styled_table(doc, ["Risiko", "Dampak", "Probabilitas", "Mitigasi"], risk_data, [1.2, 0.8, 0.8, 3.5])
    
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # BAB 10: LAMPIRAN
    # ═══════════════════════════════════════════════════════════
    
    add_section_heading(doc, "BAB 10: LAMPIRAN", level=1)
    add_separator(doc)
    
    # 10.1 API Reference
    add_section_heading(doc, "10.1 API Reference — Detail Endpoint", level=2)
    
    add_body_paragraph(doc,
        "Dokumentasi lengkap untuk setiap API endpoint. Semua endpoint menggunakan format JSON "
        "dan mengikuti standar REST API.")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Authentication API", bold=True, size=11, color=BURGUNDY)
    
    api_ref_auth = [
        ["POST /api/auth/login", 
         "Request: { email: string, password: string }\nResponse: { user: User, session: { token, expiresAt } }\nError: { error: string, code: number }",
         "200 OK, 401 Unauthorized, 429 Too Many Requests"],
        ["POST /api/auth/logout",
         "Request: (none, requires session)\nResponse: { message: 'Logged out' }",
         "200 OK"],
        ["GET /api/auth/session",
         "Request: (none)\nResponse: { user: User | null, isAuthenticated: boolean }",
         "200 OK"],
    ]
    create_styled_table(doc, ["Endpoint", "Request/Response", "Status Codes"], api_ref_auth, [1.5, 3.0, 1.8])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Pages API", bold=True, size=11, color=BURGUNDY)
    
    api_ref_pages = [
        ["GET /api/pages?page=1&limit=25&type=PGA&status=PUBLISHED",
         "Response: { data: Page[], total: number, page: number, totalPages: number }",
         "200 OK"],
        ["POST /api/pages",
         "Request: { pageNumber, type, title, content, domainId?, status: 'DRAFT' }\nResponse: { page: Page }",
         "201 Created, 400 Validation Error, 403 Forbidden"],
        ["PUT /api/pages/:id",
         "Request: { title?, content?, status?, domainId? }\nResponse: { page: Page }",
         "200 OK, 404 Not Found, 403 Forbidden"],
        ["DELETE /api/pages/:id",
         "Response: { message: 'Page deleted', pageId: string }",
         "200 OK, 404 Not Found, 403 Forbidden"],
    ]
    create_styled_table(doc, ["Endpoint", "Request/Response", "Status Codes"], api_ref_pages, [1.5, 3.0, 1.8])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line=312)
    add_formatted_run(p, "Standard Error Response Format:", bold=True, size=11, color=BURGUNDY)
    
    add_info_box(doc, "ERROR RESPONSE FORMAT",
        '{\n'
        '  "error": {\n'
        '    "code": "VALIDATION_ERROR",\n'
        '    "message": "Invalid input data",\n'
        '    "details": [\n'
        '      { "field": "email", "message": "Invalid email format" }\n'
        '    ]\n'
        '  }\n'
        '}')
    
    # 10.2 Database ERD
    add_section_heading(doc, "10.2 Database ERD — Entity Relationship Diagram", level=2)
    
    add_body_paragraph(doc,
        "Berikut adalah deskripsi hubungan antar entitas dalam database CMS KNBMP PGA-72:")
    
    erd_relationships = [
        ["users → pages", "One-to-Many", "Satu user dapat membuat banyak halaman (authorId)"],
        ["users → pages", "One-to-Many", "Satu user dapat me-review banyak halaman (reviewerId)"],
        ["users → audit_log", "One-to-Many", "Satu user memiliki banyak log aktivitas"],
        ["users → media", "One-to-Many", "Satu user dapat upload banyak media"],
        ["users → settings", "One-to-Many", "Satu user dapat mengubah banyak pengaturan"],
        ["domains → pages", "One-to-Many", "Satu domain memiliki banyak halaman PGA"],
        ["domains → pga_pillars", "One-to-Many", "Satu domain memiliki tepat 8 PGA"],
        ["pages → page_revisions", "One-to-Many", "Satu halaman memiliki banyak versi revisi"],
        ["pages → audit_log", "One-to-Many", "Satu halaman memiliki banyak log perubahan"],
        ["pga_pillars → pages", "One-to-One", "Setiap PGA memiliki satu halaman flipbook"],
    ]
    create_styled_table(doc, ["Relasi", "Kardinalitas", "Deskripsi"], erd_relationships, [1.5, 1.2, 3.6])
    
    # 10.3 Skema Prisma
    add_section_heading(doc, "10.3 Skema Prisma", level=2)
    
    add_body_paragraph(doc,
        "Berikut adalah skema Prisma lengkap untuk database CMS KNBMP PGA-72:")
    
    prisma_schema = """// prisma/schema.prisma

generator client {
  provider = "prisma-client-js"
}

datasource db {
  provider = "sqlite"
  url      = env("DATABASE_URL")
}

enum Role {
  SUPER_ADMIN
  EDITOR
  REVIEWER
}

enum PageType {
  COVER
  FOREWORD
  MUKADIMAH
  TOC
  PGA
  PHILOSOPHY
  PACT
  BACK_COVER
}

enum ContentStatus {
  DRAFT
  PENDING_REVIEW
  APPROVED
  PUBLISHED
}

enum AuditAction {
  CREATE
  UPDATE
  DELETE
  PUBLISH
  LOGIN
  LOGOUT
}

model User {
  id           String      @id @default(uuid())
  name         String
  email        String      @unique
  password     String
  role         Role        @default(EDITOR)
  avatar       String?
  isActive     Boolean     @default(true)
  lastLogin    DateTime?
  createdAt    DateTime    @default(now())
  updatedAt    DateTime    @updatedAt

  pages        Page[]      @relation("PageAuthor")
  reviews      Page[]      @relation("PageReviewer")
  media        Media[]
  auditLogs    AuditLog[]
  revisions    PageRevision[]

  @@map("users")
}

model Domain {
  id        String   @id @default(uuid())
  code      String   @unique
  name      String
  nameEn    String
  description String
  color     String
  icon      String?
  order     Int
  pgaCount  Int      @default(8)

  pages     Page[]
  pgas      PgaPillar[]

  @@map("domains")
}

model Page {
  id          String        @id @default(uuid())
  pageNumber  Int           @unique
  type        PageType
  title       String
  slug        String        @unique
  content     String        @default("")
  status      ContentStatus @default(DRAFT)
  domainId    String?
  authorId    String
  reviewerId  String?
  publishedAt DateTime?
  createdAt   DateTime      @default(now())
  updatedAt   DateTime      @updatedAt

  domain      Domain?       @relation(fields: [domainId], references: [id])
  author      User          @relation("PageAuthor", fields: [authorId], references: [id])
  reviewer    User?         @relation("PageReviewer", fields: [reviewerId], references: [id])
  revisions   PageRevision[]
  auditLogs   AuditLog[]

  @@map("pages")
}

model PgaPillar {
  id              String   @id @default(uuid())
  code            String   @unique
  name            String
  nameEn          String
  description     String
  vision          String?
  domainId        String
  orderInDomain   Int
  color           String
  dimensions      String   @default("[]")
  principles      String   @default("[]")
  crossReferences String   @default("[]")
  metadata        String   @default("{}")

  domain          Domain   @relation(fields: [domainId], references: [id])
  page            Page?

  @@map("pga_pillars")
}

model Media {
  id          String   @id @default(uuid())
  filename    String
  filepath    String
  mimeType    String
  size        Int
  alt         String?
  uploadedById String
  createdAt   DateTime @default(now())

  uploadedBy  User     @relation(fields: [uploadedById], references: [id])

  @@map("media")
}

model Setting {
  id          String   @id @default(uuid())
  key         String   @unique
  value       String
  type        String   @default("STRING")
  description String?
  updatedAt   DateTime @updatedAt
  updatedById String

  updatedBy   User     @relation(fields: [updatedById], references: [id])

  @@map("settings")
}

model AuditLog {
  id          String      @id @default(uuid())
  userId      String
  action      AuditAction
  entityType  String
  entityId    String
  oldValue    String?
  newValue    String?
  ipAddress   String?
  userAgent   String?
  createdAt   DateTime    @default(now())

  user        User        @relation(fields: [userId], references: [id])
  page        Page?       @relation(fields: [entityId], references: [id])

  @@map("audit_log")
}

model PageRevision {
  id             String   @id @default(uuid())
  pageId         String
  version        Int
  content        String
  changeSummary  String?
  createdById    String
  createdAt      DateTime @default(now())

  page           Page     @relation(fields: [pageId], references: [id])
  createdBy      User     @relation(fields: [createdById], references: [id])

  @@unique([pageId, version])
  @@map("page_revisions")
}"""

    # Add Prisma schema as code block in a table
    code_table = doc.add_table(rows=1, cols=1)
    code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = code_table.rows[0].cells[0]
    set_cell_shading(cell, "F8F6F2")
    set_cell_borders(cell, left={"sz": 6, "color": "C5A059"})
    
    # Title
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=4, after=2, line=240)
    add_formatted_run(p, "// prisma/schema.prisma", bold=False, size=8, color=MEDIUM_TEXT, font_name="Courier New")
    
    # Schema content
    for line in prisma_schema.strip().split("\n"):
        p = cell.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=200)
        add_formatted_run(p, line, size=7, color=DARK_TEXT, font_name="Courier New")
    
    # ── FOOTER (Page Numbers) ──
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = fp.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        
        run2 = fp.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        
        run3 = fp.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)
        
        # Footer text
        fp.add_run(" — ")
        run4 = fp.add_run("PRD KNBMP PGA-72 CMS v1.0 — Rahasia")
        run4.font.size = Pt(8)
        run4.font.color.rgb = MEDIUM_TEXT
        run4.font.name = "Calibri"
    
    # ── SAVE ──
    output_path = "/home/z/my-project/download/PRD-KNBMP-PGA72-CMS.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    path = create_prd()
    
    # Count pages (approximate)
    from docx import Document as Doc
    d = Doc(path)
    paragraph_count = len(d.paragraphs)
    table_count = len(d.tables)
    print(f"Total paragraphs: {paragraph_count}")
    print(f"Total tables: {table_count}")
    print("Document generation complete!")
